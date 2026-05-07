from io import BytesIO
from xml.sax.saxutils import escape

from docx import Document
from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.auth import get_current_user
from app.routers.plans import get_owned_plan
from app.utils import plan_to_document_lines


router = APIRouter(prefix="/export", tags=["export"])


@router.post("/pdf/{plan_id}")
async def export_pdf(plan_id: str, current_user: dict = Depends(get_current_user)):
    plan = await get_owned_plan(plan_id, current_user["id"])
    buffer = BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=LETTER, title="AI Business Plan Draft")
    styles = getSampleStyleSheet()
    story = []

    for level, text in plan_to_document_lines(plan):
        style = styles["Title"] if level == 0 else styles["BodyText"]
        story.append(Paragraph(escape(str(text)), style))
        story.append(Spacer(1, 8))

    pdf.build(story)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="business-plan-{plan_id}.pdf"'
        },
    )


@router.post("/docx/{plan_id}")
async def export_docx(plan_id: str, current_user: dict = Depends(get_current_user)):
    plan = await get_owned_plan(plan_id, current_user["id"])
    document = Document()

    for level, text in plan_to_document_lines(plan):
        if level == 0:
            document.add_heading(str(text), level=1)
        elif level == 1:
            document.add_paragraph(str(text))
        else:
            document.add_paragraph(str(text), style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="business-plan-{plan_id}.docx"'
        },
    )
